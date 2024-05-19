from docx import Document

from docx.enum.table import WD_ROW_HEIGHT_RULE
import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__))) 
from fanyi import fanyi
def is_float(text):
    try:
        float(text)
        return True
    except ValueError:
        return False


def replace_whitespace_with_html(s):
    # 将空格替换为&nbsp;
    s = s.replace(' ', '&nbsp;')
    # 将换行符(\n)替换为<br>以在HTML中实现换行
    s = s.replace('\n', '<br>')
    return s
def do_text(text,is_count_mode):
    if text in ['\n', '']:
        return text
    if text.isnumeric() or is_float(text):
        return text
    translated_text = fanyi(text)
    return translated_text    
        
def do_paragraph(paragraph,is_count_mode):
   
    # 检查段落中的每一个运行元素
    for run in paragraph.runs:
        # 如果运行元素包含图片或数学公式，跳过这个段落
        if 'graphicData' in run._r.xml or 'oMathPara' in run._r.xml:
            return

    # 处理段落的文本
    
    processed_text = do_text(paragraph.text,is_count_mode)
     
    
    # 清空原有的段落
    paragraph.clear()
    # 添加处理后的文本到原有段落
    paragraph.add_run(processed_text)
def do_table(table,is_count_mode):
    for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO

            # 遍历行中的所有单元格
            for cell in row.cells:
                # 遍历单元格中的所有段落
                for paragraph in cell.paragraphs:
                    do_paragraph(paragraph, is_count_mode)
                    # paragraph.paragraph_format.space_after = Cm(0)  # 移除段落后的间距

                # 检查嵌套表格是否存在
                if cell.tables is not None:
                    # 递归处理嵌套表格
                    for nested_table in cell.tables:
                        do_table(nested_table, is_count_mode)
def docx_fanyi(doc_file,doc_file_zh):
    doc = Document(doc_file)
    # 处理table
    
    for table in doc.tables:
        do_table(table,True)

    if(doc.paragraphs is not None):
        for paragraph in doc.paragraphs:
            do_paragraph(paragraph,True)

    doc.save(doc_file_zh)